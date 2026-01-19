"""
Document processing for Environmental Semantic Search Tool.

Handles loading, parsing, and chunking various document formats.
"""

import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter
from loguru import logger

from .config import settings
from .utils import (
    Timer,
    clean_text,
    get_file_extension,
    get_file_hash,
    is_supported_file,
    list_documents,
    DocumentError,
)


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class Document:
    """Represents a document or document chunk."""
    
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    @property
    def id(self) -> str:
        """Generate unique ID based on content hash."""
        return hashlib.md5(self.content.encode()).hexdigest()[:16]
    
    def __len__(self) -> int:
        return len(self.content)
    
    def __repr__(self) -> str:
        source = self.metadata.get("source", "unknown")
        return f"Document(source={source}, length={len(self)})"


@dataclass 
class ProcessingResult:
    """Result of document processing."""
    
    documents: List[Document]
    total_chunks: int
    total_characters: int
    sources: List[str]
    errors: List[str] = field(default_factory=list)


# =============================================================================
# Document Loaders
# =============================================================================

def load_pdf(file_path: Path) -> str:
    """Load text from PDF file."""
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(str(file_path))
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise DocumentError(f"Error loading PDF {file_path}: {e}")


def load_docx(file_path: Path) -> str:
    """Load text from DOCX file."""
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise DocumentError(f"Error loading DOCX {file_path}: {e}")


def load_text(file_path: Path) -> str:
    """Load text from plain text file."""
    try:
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise DocumentError(f"Could not decode {file_path} with any encoding")
    except Exception as e:
        raise DocumentError(f"Error loading text file {file_path}: {e}")


def load_html(file_path: Path) -> str:
    """Load text from HTML file."""
    try:
        from bs4 import BeautifulSoup
        
        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
        
        # Remove script and style elements
        for element in soup(["script", "style", "nav", "header", "footer"]):
            element.decompose()
        
        return soup.get_text(separator="\n")
    except Exception as e:
        raise DocumentError(f"Error loading HTML {file_path}: {e}")


def load_markdown(file_path: Path) -> str:
    """Load text from Markdown file."""
    return load_text(file_path)


# Loader mapping
LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".doc": load_docx,
    ".txt": load_text,
    ".md": load_markdown,
    ".html": load_html,
    ".htm": load_html,
}


# =============================================================================
# Document Processor
# =============================================================================

class DocumentProcessor:
    """
    Process documents: load, clean, and chunk.
    """
    
    def __init__(
        self,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
        min_chunk_size: Optional[int] = None,
    ):
        """
        Initialize document processor.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
            min_chunk_size: Minimum chunk size to keep
        """
        self.chunk_size = chunk_size or settings.chunk_size
        self.chunk_overlap = chunk_overlap or settings.chunk_overlap
        self.min_chunk_size = min_chunk_size or settings.min_chunk_size
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        
        logger.info(
            f"DocumentProcessor initialized: "
            f"chunk_size={self.chunk_size}, overlap={self.chunk_overlap}"
        )
    
    def load_file(self, file_path: Path) -> str:
        """
        Load content from a file.
        
        Args:
            file_path: Path to file
            
        Returns:
            File content as string
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise DocumentError(f"File not found: {file_path}")
        
        if not is_supported_file(file_path):
            raise DocumentError(f"Unsupported file type: {file_path.suffix}")
        
        extension = get_file_extension(file_path)
        loader = LOADERS.get(extension)
        
        if not loader:
            raise DocumentError(f"No loader for: {extension}")
        
        with Timer(f"Loading {file_path.name}"):
            content = loader(file_path)
        
        return content
    
    def chunk_text(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Document]:
        """
        Split text into chunks.
        
        Args:
            text: Text to split
            metadata: Metadata to add to each chunk
            
        Returns:
            List of Document chunks
        """
        if not text.strip():
            return []
        
        metadata = metadata or {}
        
        # Split text
        chunks = self.text_splitter.split_text(text)
        
        # Create documents
        documents = []
        for i, chunk in enumerate(chunks):
            # Clean chunk
            cleaned = clean_text(chunk)
            
            # Skip if too short
            if len(cleaned) < self.min_chunk_size:
                continue
            
            # Create document with metadata
            doc = Document(
                content=cleaned,
                metadata={
                    **metadata,
                    "chunk_index": i,
                    "chunk_size": len(cleaned),
                }
            )
            documents.append(doc)
        
        return documents
    
    def process_file(self, file_path: Path) -> List[Document]:
        """
        Load and chunk a single file.
        
        Args:
            file_path: Path to file
            
        Returns:
            List of Document chunks
        """
        file_path = Path(file_path)
        
        # Load content
        content = self.load_file(file_path)
        
        # Create metadata
        metadata = {
            "source": file_path.name,
            "source_path": str(file_path),
            "file_type": get_file_extension(file_path),
            "file_hash": get_file_hash(file_path),
        }
        
        # Chunk
        documents = self.chunk_text(content, metadata)
        
        logger.info(f"Processed {file_path.name}: {len(documents)} chunks")
        
        return documents
    
    def process_directory(
        self,
        directory: Path,
        recursive: bool = True,
        show_progress: bool = True,
    ) -> ProcessingResult:
        """
        Process all documents in a directory.
        
        Args:
            directory: Directory path
            recursive: Search recursively
            show_progress: Show progress logs
            
        Returns:
            ProcessingResult with all documents
        """
        directory = Path(directory)
        
        if not directory.exists():
            raise DocumentError(f"Directory not found: {directory}")
        
        # Find all documents
        files = list_documents(directory, recursive=recursive)
        logger.info(f"Found {len(files)} documents in {directory}")
        
        all_documents = []
        sources = []
        errors = []
        
        for i, file_path in enumerate(files):
            if show_progress:
                logger.info(f"Processing [{i+1}/{len(files)}]: {file_path.name}")
            
            try:
                documents = self.process_file(file_path)
                all_documents.extend(documents)
                sources.append(str(file_path))
            except DocumentError as e:
                logger.warning(f"Skipping {file_path.name}: {e}")
                errors.append(str(e))
        
        result = ProcessingResult(
            documents=all_documents,
            total_chunks=len(all_documents),
            total_characters=sum(len(d) for d in all_documents),
            sources=sources,
            errors=errors,
        )
        
        logger.info(
            f"Processing complete: {result.total_chunks} chunks "
            f"from {len(result.sources)} files"
        )
        
        return result
    
    def process_texts(
        self,
        texts: List[str],
        sources: Optional[List[str]] = None,
    ) -> List[Document]:
        """
        Process a list of text strings.
        
        Args:
            texts: List of texts
            sources: Optional source names
            
        Returns:
            List of Document chunks
        """
        if sources is None:
            sources = [f"text_{i}" for i in range(len(texts))]
        
        all_documents = []
        
        for text, source in zip(texts, sources):
            metadata = {"source": source}
            documents = self.chunk_text(text, metadata)
            all_documents.extend(documents)
        
        return all_documents


# =============================================================================
# Utility Functions
# =============================================================================

def deduplicate_documents(documents: List[Document]) -> List[Document]:
    """
    Remove duplicate documents based on content.
    
    Args:
        documents: List of documents
        
    Returns:
        Deduplicated list
    """
    seen_hashes = set()
    unique = []
    
    for doc in documents:
        doc_hash = doc.id
        if doc_hash not in seen_hashes:
            seen_hashes.add(doc_hash)
            unique.append(doc)
    
    removed = len(documents) - len(unique)
    if removed > 0:
        logger.info(f"Removed {removed} duplicate documents")
    
    return unique


def get_document_stats(documents: List[Document]) -> Dict[str, Any]:
    """
    Get statistics about documents.
    
    Args:
        documents: List of documents
        
    Returns:
        Statistics dictionary
    """
    if not documents:
        return {"count": 0}
    
    lengths = [len(d) for d in documents]
    sources = set(d.metadata.get("source", "unknown") for d in documents)
    
    return {
        "count": len(documents),
        "total_characters": sum(lengths),
        "avg_length": sum(lengths) / len(lengths),
        "min_length": min(lengths),
        "max_length": max(lengths),
        "unique_sources": len(sources),
    }
