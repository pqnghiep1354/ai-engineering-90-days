"""
DOCX Generator for EIA Reports.
"""

from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from loguru import logger

from ..config import EIAReport, EIASection, ProjectInput


class DocxGenerator:
    """Generate professional DOCX reports from EIA data."""
    
    def __init__(self, template_path: Optional[str] = None):
        """Initialize generator with optional template."""
        self.template_path = template_path
    
    def generate(
        self,
        report: EIAReport,
        output_path: str,
    ) -> str:
        """
        Generate DOCX report.
        
        Args:
            report: EIA report data
            output_path: Output file path
            
        Returns:
            Path to generated file
        """
        logger.info(f"Generating DOCX report: {output_path}")
        
        # Create document
        if self.template_path and Path(self.template_path).exists():
            doc = Document(self.template_path)
        else:
            doc = Document()
            self._setup_styles(doc)
        
        # Add content
        self._add_cover_page(doc, report.project)
        self._add_table_of_contents(doc)
        self._add_executive_summary(doc, report.executive_summary)
        
        for section in report.sections:
            self._add_section(doc, section)
        
        self._add_appendices(doc, report.appendices)
        self._add_footer(doc, report)
        
        # Save
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        logger.info(f"Report saved to: {output_path}")
        return str(output_path)
    
    def _setup_styles(self, doc: Document) -> None:
        """Setup document styles."""
        styles = doc.styles
        
        # Title style
        if 'EIA Title' not in [s.name for s in styles]:
            title_style = styles.add_style('EIA Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(18)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Heading 1
        h1 = styles['Heading 1']
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(6)
        
        # Heading 2
        h2 = styles['Heading 2']
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.paragraph_format.space_before = Pt(10)
        h2.paragraph_format.space_after = Pt(4)
        
        # Normal text
        normal = styles['Normal']
        normal.font.size = Pt(11)
        normal.font.name = 'Times New Roman'
        normal.paragraph_format.line_spacing = 1.5
    
    def _add_cover_page(self, doc: Document, project: ProjectInput) -> None:
        """Add cover page."""
        # Add spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Ministry header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("BỘ TÀI NGUYÊN VÀ MÔI TRƯỜNG")
        run.bold = True
        run.font.size = Pt(14)
        
        doc.add_paragraph()
        
        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("BÁO CÁO")
        run.bold = True
        run.font.size = Pt(16)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ĐÁNH GIÁ TÁC ĐỘNG MÔI TRƯỜNG")
        run.bold = True
        run.font.size = Pt(18)
        
        doc.add_paragraph()
        
        # Project name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"DỰ ÁN: {project.name.upper()}")
        run.bold = True
        run.font.size = Pt(14)
        
        # Add spacing
        for _ in range(5):
            doc.add_paragraph()
        
        # Investor info
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"CHỦ DỰ ÁN: {project.investor_name or 'N/A'}")
        
        # Location
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"ĐỊA ĐIỂM: {project.location}")
        
        # Page break
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document) -> None:
        """Add table of contents."""
        doc.add_heading("MỤC LỤC", level=1)
        
        # Placeholder for TOC
        p = doc.add_paragraph()
        p.add_run("[Mục lục sẽ được cập nhật tự động]")
        p.italic = True
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, summary: str) -> None:
        """Add executive summary."""
        doc.add_heading("TÓM TẮT BÁO CÁO", level=1)
        
        # Parse and add summary content
        for line in summary.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                # Handle bold text
                if '**' in line:
                    parts = line[2:].split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
                else:
                    p.add_run(line[2:])
            else:
                doc.add_paragraph(line)
        
        doc.add_page_break()
    
    def _add_section(self, doc: Document, section: EIASection) -> None:
        """Add a report section."""
        # Section heading
        doc.add_heading(f"CHƯƠNG {section.id}: {section.title}", level=1)
        
        # English title
        if section.title_en:
            p = doc.add_paragraph()
            run = p.add_run(f"({section.title_en})")
            run.italic = True
            run.font.size = Pt(10)
        
        # Section content
        if section.content:
            self._add_formatted_content(doc, section.content)
        
        # Subsections
        for subsection in section.subsections:
            doc.add_heading(f"{subsection.id}. {subsection.title}", level=2)
            if subsection.content:
                self._add_formatted_content(doc, subsection.content)
        
        # Tables
        for table_data in section.tables:
            self._add_table(doc, table_data)
        
        doc.add_page_break()
    
    def _add_formatted_content(self, doc: Document, content: str) -> None:
        """Add formatted content (parse markdown-like syntax)."""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Headings
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])
            elif line[0].isdigit() and line[1] in '.):':
                p = doc.add_paragraph(style='List Number')
                p.add_run(line[3:].strip())
            # Normal paragraph
            else:
                doc.add_paragraph(line)
    
    def _add_table(
        self,
        doc: Document,
        table_data: Dict[str, Any],
    ) -> None:
        """Add a table to document."""
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        title = table_data.get("title", "")
        
        if title:
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.bold = True
        
        if not headers or not rows:
            return
        
        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = str(cell_data)
        
        doc.add_paragraph()
    
    def _add_appendices(
        self,
        doc: Document,
        appendices: List[Dict[str, Any]],
    ) -> None:
        """Add appendices."""
        if not appendices:
            return
        
        doc.add_heading("PHỤ LỤC", level=1)
        
        for i, appendix in enumerate(appendices, 1):
            title = appendix.get("title", f"Phụ lục {i}")
            content = appendix.get("content", "")
            
            doc.add_heading(f"Phụ lục {i}: {title}", level=2)
            if content:
                doc.add_paragraph(content)
    
    def _add_footer(self, doc: Document, report: EIAReport) -> None:
        """Add footer with metadata."""
        doc.add_paragraph()
        doc.add_paragraph("---")
        
        p = doc.add_paragraph()
        p.add_run(f"Báo cáo được tạo: {report.generated_at}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if report.compliance_score > 0:
            p = doc.add_paragraph()
            p.add_run(f"Điểm đánh giá: {report.compliance_score:.1f}/100")
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
